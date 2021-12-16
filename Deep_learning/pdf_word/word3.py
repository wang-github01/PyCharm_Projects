#coding=utf-8

import fitz # pdf转为图片
from aip import AipOcr # 图片文字识别
import time # 程序运行时间间隔以避免出错
import docx # 将识别结果保存为docx文件
from docx.oxml.ns import qn # 设置docx文件的字体

""" 你的 APPID AK SK """
APP_ID = '25230547'
API_KEY = 'zgPNf0tnh89Du22XHyuiHI7K'
SECRET_KEY = 'YCZIi0ouwFeRXq8lSI3TdBGmhS3srG8V'

client = AipOcr(APP_ID, API_KEY, SECRET_KEY)

'''
将PDF转化为图片
pdfPath pdf文件的路径
imgPath 图像要保存的路径
zoom_x x方向的缩放系数
zoom_y y方向的缩放系数
rotation_angle 旋转角度
zoom_x和zoom_y一般取相同值，值越大，图像分辨率越高
返回目标pdf的名称和页数，便于下一步操作

'''
def pdf_image(pdfPath, imgPath, zoom_x=2.0, zoom_y=2.0, rotation_angle=0):
    # 获取pdf文件名称
    name = pdfPath.split("\\")[-1].split('.pdf')[0]
    # 打开PDF文件
    pdf = fitz.open(pdfPath)
    # 获取pdf页数
    num = pdf.pageCount
    # 逐页读取PDF
    for pg in range(0, num):
        page = pdf[pg]
        # 设置缩放和旋转系数
        trans = fitz.Matrix(2.0, 2.0).preRotate(0)
        pm = page.getPixmap(matrix=trans, alpha=False)
        # 开始写图像
        pm.writePNG(imgPath + name + "_" + str(pg) + ".png")
    pdf.close()
    return name, num
'''
将图片读取为docx文件
imgPath 图像所在路径
生成的docx也保存在图像所在路径中
name为pdf名称（不含后缀）
num为pdf页数
name和num均可由上一个函数返回

'''
def ReadDetail_docx(imgPath, name, num):
    # 建立一个空doc文档
    doc = docx.Document()
    # 设置全局字体
    doc.styles["Normal"].font.name=u"宋体"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 读取图片
    j = 0
    for n in range(0,num):
        i = open(imgPath+name+"_"+str(n)+".png",'rb')
        time.sleep(0.1)
        img = i.read()
        #message = client.office(img)
        message = client.accurate(img)
        print(j)
        print(message)
        j = j+1

        content = message.get('words_result')
        #content = message.get('results')
        print(content)
        # 将内容写入doc文档
        #doc.add_paragraph(content.get('words'))
        for i in range(len(content)):
            doc.add_paragraph(content[i].get('words'))
            print(content[i].get("words"))
            print(content[i].get('location'))
        print("-----------")
    # 保存doc文档
    doc.save(imgPath + name + '.docx')
def pdf_to_docx(pdfPath, imgPath, zoom_x=5, zoom_y=5, rotation_angle=0):
    print("正在将pdf文件转换为图片...")
    # 调用函数一将pdf转换为图片，并获得文件名和页数
    name_, num_ = pdf_image(pdfPath, imgPath, zoom_x, zoom_y, rotation_angle)
    print("转换成功！")
    print("正在读取图片内容...")
    # 调用函数二逐页读取图片并逐行保存在docx文件中
    ReadDetail_docx(imgPath, name_, num_)
    print("名为 {}.pdf 的pdf文件共有{}页，已成功转换为docx文件！".format(name_, num_))
# pdf储存路径
pdf_path = r"E:pdf\2021.pdf"
# 图片和生成的docx文件的储存路径
img_path = "E:pdf\image\\"
# 调用函数
pdf_to_docx(pdf_path, img_path)